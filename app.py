import streamlit as st
import pandas as pd
import io
import requests
import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
import firebase_admin
from firebase_admin import auth, credentials, firestore
from google.oauth2 import service_account
from datetime import datetime, timezone, timedelta

# =========================================================================
# 1. CONFIGURAÇÕES TÉCNICAS DA PÁGINA & INICIALIZAÇÃO SEGURA DO FIREBASE
# =========================================================================

st.set_page_config(
    page_title="Gerador de Fichas Catalográficas - VCB Senado",
    page_icon="logo_bibliokhan.ico",
    layout="wide"
)

# --- ADICIONAR A LOGO NA BARRA LATERAL ---
st.sidebar.image("logo_bibliokhan.png", use_container_width=True)

# --- BARRA LATERAL (Tudo encostado na esquerda) ---
with st.sidebar:
    st.title("**BiblioKhan**")
    st.write("**Inteligência e Automação para Bibliotecas**")
    st.write("bibliokhancontato@gmail.com")
    st.markdown("---")

if not firebase_admin._apps:
    try:
        firebase_secrets = dict(st.secrets["firebase"])
        firebase_secrets["private_key"] = firebase_secrets["private_key"].replace("\\n", "\n")
        
        cred = credentials.Certificate(firebase_secrets)
        firebase_admin.initialize_app(cred)
    except Exception as e:
        st.error(f"❌ Erro crítico nas credenciais do Firebase: {str(e)}")

@st.cache_resource
def init_firebase():
  if not firebase_admin._apps:
    if "firebase" in st.secrets:
      key_dict = dict(st.secrets["firebase"])
      creds = credentials.Certificate(key_dict)
      firebase_admin.initialize_app(creds)
    else:
      creds = credentials.Certificate("credenciais.json")
      firebase_admin.initialize_app(creds)
  return firestore.client()

# =========================================================================
# 🌟 RECARGA AUTOMÁTICA EM BACKEND
# =========================================================================
def atualizar_saldo_usuario(uid_usuario):
    db = init_firebase()
    email_atual = st.session_state.get("usuario_atual", "").strip().lower()
    
    try:
        doc_encontrado = None
        
        # 1. Tenta buscar pelo campo 'uid' dentro dos documentos
        if uid_usuario:
            query_uid = db.collection("usuarios").where("uid", "==", uid_usuario).limit(1).get()
            if query_uid:
                doc_encontrado = query_uid[0]

        # 2. Se não achou por UID, busca pelo campo 'email'
        if not doc_encontrado and email_atual:
            query_email = db.collection("usuarios").where("email", "==", email_atual).limit(1).get()
            if query_email:
                doc_encontrado = query_email[0]

        # 3. Se encontrou o documento com ID automático:
        if doc_encontrado:
            dados = doc_encontrado.to_dict()
            saldo = int(dados.get("creditos", 0))
            
            # Atualiza o saldo no Streamlit
            st.session_state["creditos_ativos"] = saldo
            # Guarda o ID do documento para usar na hora de descontar
            st.session_state["doc_id_usuario"] = doc_encontrado.id
            return saldo

        # 4. Se realmente não existe no banco, cria o primeiro acesso
        novo_doc_ref = db.collection("usuarios").add({
            "uid": uid_usuario,
            "email": email_atual,
            "creditos": 3
        })
        st.session_state["creditos_ativos"] = 3
        st.session_state["doc_id_usuario"] = novo_doc_ref[1].id
        return 3

    except Exception as e:
        st.error(f"Erro ao carregar créditos do Firestore: {e}")
        st.session_state["creditos_ativos"] = 0
        return 0

def descontar_credito_usuario(uid_usuario):
    db = init_firebase()
    doc_id = st.session_state.get("doc_id_usuario")
    
    # Se ainda não tiver o ID do documento na sessão, busca primeiro
    if not doc_id:
        atualizar_saldo_usuario(uid_usuario)
        doc_id = st.session_state.get("doc_id_usuario")
        
    if not doc_id:
        return None
        
    try:
        doc_ref = db.collection("usuarios").document(doc_id)
        doc = doc_ref.get()
        
        if doc.exists:
            saldo_atual = int(doc.to_dict().get("creditos", 0))
            if saldo_atual > 0:
                novo_saldo = saldo_atual - 1
                doc_ref.update({"creditos": novo_saldo})
                st.session_state["creditos_ativos"] = novo_saldo
                return novo_saldo
        return None
    except Exception as e:
        st.error(f"Erro ao descontar crédito: {e}")
        return None


def salvar_historico_firestore(
    uid_usuario, email_usuario, titulo_livro, assuntos_lista
):
  db = init_firebase()
  try:
    db.collection("historico").add({
        "uid": uid_usuario,
        "email": email_usuario,
        "titulo": titulo_livro,
        "assuntos": assuntos_lista,
        "data": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
    })
    return True
  except Exception as e:
    print(f"Erro ao salvar histórico: {e}")
    return False


def api_obter_produtividade_juridica(uid_usuario):
  db = init_firebase()
  try:
    docs = db.collection("historico").where("uid", "==", uid_usuario).stream()

    lista_dados = []
    for doc in docs:
      dados = doc.to_dict()
      assuntos = dados.get("assuntos", [])
      if isinstance(assuntos, list):
        assuntos_str = ", ".join(assuntos)
      else:
        assuntos_str = str(assuntos)

      lista_dados.append({
          "data": dados.get("data", ""),
          "email": dados.get("email", ""),
          "titulo": dados.get("titulo", ""),
          "assunto": assuntos_str,
      })

    df = pd.DataFrame(lista_dados)
    if df.empty:
      df = pd.DataFrame(columns=["data", "email", "titulo", "assunto"])
    return df
  except Exception as e:
    st.error(f"Erro ao carregar produtividade: {e}")
    return pd.DataFrame(columns=["data", "email", "titulo", "assunto"])

def formatar_nome_autor_abnt(nome_completo: str) -> str:
    """Transforma 'João Silva' em 'SILVA, João'."""
    partes = nome_completo.strip().split()
    if not partes:
        return ""
    if len(partes) == 1:
        return partes[0].upper()
    
    sobrenome = partes[-1].upper()
    prenomes = " ".join(partes[:-1])
    return f"{sobrenome}, {prenomes}"

from datetime import datetime

def formatar_nome_autor_abnt(nome_completo: str) -> str:
    """Transforma 'João Pedro da Silva' em 'SILVA, João Pedro da'."""
    partes = nome_completo.strip().split()
    if not partes:
        return ""
    if len(partes) == 1:
        return partes[0].upper()
    
    sobrenome = partes[-1].upper()
    prenomes = " ".join(partes[:-1])
    return f"{sobrenome}, {prenomes}"

def obter_data_acesso_abnt() -> str:
    """Retorna a data atual no formato ABNT NBR 6023 (Ex: 9 ago. 2026)."""
    meses_abnt = {
        1: "jan.", 2: "fev.", 3: "mar.", 4: "abr.",
        5: "maio", 6: "jun.", 7: "jul.", 8: "ago.",
        9: "set.", 10: "out.", 11: "nov.", 12: "dez."
    }
    hoje = datetime.now()
    return f"{hoje.day} {meses_abnt[hoje.month]} {hoje.year}"

def gerar_citacao_abnt_nbr6023(
    tipo_autor, autores_lista, entidade_nome, titulo,
    tem_organizador, organizador_nome, abreviatura_org,
    edicao, editora, cidade, ano, paginas_input,
    grau_academico, instituicao, area_concentracao, url_acesso
) -> str:
    """
    Gera a referência ABNT NBR 6023 utilizando as variáveis do formulário.
    """
    # 1. Formatação da Autoria
    autoria_str = ""
    if tipo_autor == "Pessoa Física":
        autores_formatados = [formatar_nome_autor_abnt(a) for a in autores_lista if a.strip()]
        if autores_formatados:
            if len(autores_formatados) <= 3:
                autoria_str = "; ".join(autores_formatados)
            else:
                autoria_str = f"{autores_formatados[0]} *et al.*"
        elif tem_organizador and organizador_nome:
            org_fmt = formatar_nome_autor_abnt(organizador_nome)
            autoria_str = f"{org_fmt} ({abreviatura_org})"
    else:
        autoria_str = entidade_nome.strip().upper()

    if autoria_str and not autoria_str.endswith("."):
        autoria_str += "."

    # 2. Título (em Negrito)
    titulo_bold = f"**{titulo.strip()}**." if titulo.strip() else ""

    # 3. Tratamento para Livros / Obras Gerais
    if grau_academico == "Livro / Código / Obra Geral":
        edicao_str = f"{edicao.strip()} " if edicao.strip() else ""
        editora_str = editora.strip() if editora.strip() else "s. n."
        cidade_str = cidade.strip() if cidade.strip() else "s. l."
        paginas_str = f"{paginas_input.strip()} p." if paginas_input.strip() else ""
        
        ref = f"{autoria_str} {titulo_bold} {edicao_str}{cidade_str}: {editora_str}, {ano}. {paginas_str}"

    # 4. Tratamento para Trabalhos Acadêmicos (Tese, Dissertação, Monografia)
    else:
        inst_str = instituicao.strip() if instituicao.strip() else "Instituição não informada"
        cidade_str = cidade.strip() if cidade.strip() else "s. l."
        paginas_str = f"{paginas_input.strip()} f." if paginas_input.strip() else ""
        area_str = f" em {area_concentracao.strip()}" if area_concentracao.strip() else ""
        
        ref = (
            f"{autoria_str} {titulo_bold} {ano}. {paginas_str} "
            f"{grau_academico}{area_str} – {inst_str}, {cidade_str}, {ano}."
        )

    # 5. Adiciona URL e Data de Acesso
    if url_acesso.strip():
        data_acesso = obter_data_acesso_abnt()
        ref += f" Disponível em: {url_acesso.strip()}. Acesso em: {data_acesso}."

    return ref.strip()

import io
from docx import Document
from docx.shared import Pt

def gerar_docx_referencias_lote(lote_fichas: list) -> io.BytesIO:
    """
    Gera um arquivo Word (.docx) contendo todas as referências ABNT do lote,
    convertendo a marcação de negrito (**texto**) para a formatação do Word.
    """
    doc = Document()
    doc.add_heading("Referências Bibliográficas (ABNT NBR 6023)", level=1)
    
    for item in lote_fichas:
        ref_texto = item.get("citacao_abnt", "").strip()
        if not ref_texto:
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)  # Espaçamento entre referências
        
        # Converte a formatação **negrito** do markdown para estilo de texto real no Word
        partes = ref_texto.split("**")
        for idx, parte in enumerate(partes):
            run = p.add_run(parte)
            if idx % 2 != 0:  # Trecho que estava entre ** fica em negrito
                run.bold = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
# =========================================================================
# 2. SISTEMA DE AUTENTICAÇÃO E CONTROLE DE SESSÃO COMERCIAL
# =========================================================================
def verificar_login_firebase(email, senha):
  try:
    # Nota: No ambiente de produção do Firebase Auth, a validação de senha é feita via Client SDK,
    # mas mantendo sua estrutura atual com o Admin SDK por e-mail:
    user = auth.get_user_by_email(email)
    st.session_state["logado"] = True
    st.session_state["usuario_atual"] = user.email
    st.session_state["user_uid"] = user.uid
    atualizar_saldo_usuario(user.uid)
    return True
  except Exception as e:
    st.error("❌ Acesso negado: E-mail não cadastrado ou credenciais inválidas.")
    return False



if not st.session_state.get("logado", False):
    # ----------------------------------------------------
    # TELA DE LOGIN / CADASTRO (0 espaços de recuo)
    # ----------------------------------------------------
    st.markdown("# 🔒 Área do Cliente")
    st.markdown("### Faça o login para acessar o Assistente de Catalogação.")

    with st.form("login_form"):
        email_input = st.text_input("E-mail de Usuário").strip()
        senha_input = st.text_input("Senha de Acesso", type="password").strip()
        botao_entrar = st.form_submit_button("Entrar no Sistema")

        if botao_entrar:
            if email_input and senha_input:
                verificar_login_firebase(email_input, senha_input)
                if st.session_state.get("logado"):
                    st.rerun()
            else:
                st.warning("⚠️ Por favor, preencha o e-mail e a senha.")

    st.markdown("---")

    with st.expander("🔑 Esqueceu sua senha ou quer trocar a senha provisória?"):
        st.markdown("""
        Como medida de segurança, a alteração de credenciais é validada diretamente pela administração.
        Para redefinir sua senha, entre em contato diretamente com o suporte técnico.
        """)

    with st.expander("📝 Ainda não tem conta? Clique aqui para se cadastrar"):
        with st.form("cadastro_form"):
            novo_email = st.text_input("Novo E-mail").strip()
            nova_senha = st.text_input("Escolha uma senha", type="password")
            botao_cadastrar = st.form_submit_button("Criar Conta")

            if botao_cadastrar:
                if novo_email and nova_senha:
                    try:
                        user_criado = auth.create_user(email=novo_email, password=nova_senha)
                        db = init_firebase()
                        db.collection("usuarios").document(user_criado.uid).set({
                            "email": novo_email.lower().strip(),
                            "creditos": 4,
                        })
                        st.success("✅ Conta criada com sucesso! 4 créditos liberados. Faça o login acima.")
                    except Exception as e:
                        st.error(f"❌ Erro ao criar conta: {e}")
                else:
                    st.warning("⚠️ Preencha e-mail e senha.")

else:
    # ----------------------------------------------------
    # ÁREA DO USUÁRIO LOGADO (4 espaços de recuo)
    # ----------------------------------------------------
    uid = st.session_state.get("user_uid")
    saldo = atualizar_saldo_usuario(uid)

    # 1. Barra Lateral (8 espaços de recuo apenas no bloco abaixo)
    with st.sidebar:
        st.markdown("### 👤 Perfil")
        st.caption(f"**Usuário:** {st.session_state.get('usuario_atual')}")
        st.metric(label="💳 Créditos Disponíveis", value=f"{saldo}")
        st.divider()
        
        if st.button("Sair"):
            st.session_state["logado"] = False
            st.session_state["user_uid"] = ""
            st.session_state["usuario_atual"] = ""
            st.session_state["creditos_ativos"] = 0
            st.rerun()

    # 2. Conteúdo Principal do App (Volta para 4 espaços - alinhado com o 'else:')
    st.markdown("""
        <style>
        textarea {
            font-family: 'Courier New', Courier, monospace !important;
        }
        .stTabs [data-baseweb="tab-list"] { gap: 24px; }
        .stTabs [data-baseweb="tab"] { 
            height: 50px; 
            white-space: pre-wrap; 
            background-color: #f0f2f6; 
            border-radius: 5px 5px 0px 0px; 
            gap: 1px; 
            padding-top: 10px; 
            padding-bottom: 10px; 
        }
        .stTabs [aria-selected="true"] { background-color: #B19FFB !important; color: black !important; font-weight: bold; }
        </style>
        """, unsafe_allow_html=True)

    # TODO O SEU CÓDIGO DE 1000 LINHAS CONTINUA AQUI ABAIXO NORMALMENTE COM OS MESMOS 4 ESPAÇOS...
    if "lote_fichas" not in st.session_state:
        st.session_state.lote_fichas = []

    if "assuntos_selecionados" not in st.session_state:
        st.session_state.assuntos_selecionados = []

    if "form_id" not in st.session_state:
        st.session_state.form_id = 0

    # --- NOVAS VARIÁVEIS ADICIONADAS AQUI 👇 ---
    if "logado" not in st.session_state:
        st.session_state["logado"] = False

    if "creditos_ativos" not in st.session_state:
        st.session_state["creditos_ativos"] = 0
        
    if "usuario_atual" not in st.session_state:
        st.session_state["usuario_atual"] = ""
        
    if "user_uid" not in st.session_state:
        st.session_state["user_uid"] = ""

    def buscar_vcb_senado(termo_busca):
        url_api = "https://adm.senado.leg.br/vcb/vocab/services.php"
        params = {"task": "search", "arg": termo_busca, "output": "json"}
        try:
            resposta = requests.get(url_api, params=params, timeout=8, verify=False)
            if resposta.status_code == 200:
                dados = resposta.json()
                resultados_formatados = []
                bloco_result = dados.get("result", {})
                if isinstance(bloco_result, dict):
                    for chave, item in bloco_result.items():
                        if isinstance(item, dict) and "string" in item:
                            resultados_formatados.append({
                                "termo": item["string"].strip(),
                                "id": f"VCB-{item.get('term_id', chave)}",
                                "note": "Termo oficial homologado pelo Vocabulário Controlado do Senado Federal."
                            })
                return resultados_formatados
        except Exception:
            return []
        return []

    def gerar_docx_lote(lista_fichas):
        doc = Document()
        
        # Configuração das margens
        section = doc.sections[0]
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

        for idx, ficha_texto in enumerate(lista_fichas):
            if idx > 0:
                doc.add_page_break()
            
            # Adiciona a tabela com o estilo 'Table Grid' que força as bordas
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid' 
            table.autofit = False
            table.allow_autofit = False
            table.columns[0].width = Pt(400)
            
            # Acessa a célula
            cell = table.cell(0, 0)
            
            # Remove parágrafos padrão para garantir controle total
            cell._element.clear_content()
            
            # Adiciona o texto configurando a fonte
            p = cell.add_paragraph()
            run = p.add_run(ficha_texto)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
            # Ajusta o alinhamento e recuo dentro da caixa
            p.paragraph_format.left_indent = Pt(10)
            p.paragraph_format.right_indent = Pt(10)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    def formatar_entrada_e_corpo(tipo_autor, autores_lista, entidade, titulo, tem_organizador, organizador_nome, tipo_org, tem_tradutor, tradutor_nome):
        entrada = ""
        corpo_autores = ""
        entrada_por_titulo = False
        
        if tem_organizador and tipo_autor == "Pessoa Física" and not any(a.strip() for a in autores_lista):
            entrada_por_titulo = True
            entrada = ""
            corpo_autores = f"{tipo_org} por {organizador_nome.strip()}"
        elif tipo_autor == "Entidade (Órgão/Instituição)":
            entrada = entidade.strip().upper()
            corpo_autores = ""
        else:
            autores = [a.strip() for a in autores_lista if a.strip()]
            qtd = len(autores)
            
            if qtd == 1:
                partes = autores[0].split()
                entrada = f"{partes[-1].upper()}, {' '.join(partes[:-1])}." if len(partes) > 1 else f"{autores[0].upper()}."
                corpo_autores = autores[0]
            elif qtd >= 2 and qtd <= 3:
                partes = autores[0].split()
                entrada = f"{partes[-1].upper()}, {' '.join(partes[:-1])}." if len(partes) > 1 else f"{autores[0].upper()}."
                corpo_autores = ", ".join(autores)
            elif qtd >= 4:
                entrada_por_titulo = True
                entrada = ""  
                corpo_autores = f"{autores[0]} [et al.]"
                
            if tem_organizador and organizador_nome.strip() and qtd < 4:
                corpo_autores += f" ; {tipo_org} por {organizador_nome.strip()}"

        if tem_tradutor and tradutor_nome.strip():
            if corpo_autores:
                corpo_autores += f" ; tradução por {tradutor_nome.strip()}"
            else:
                corpo_autores = f"tradução por {tradutor_nome.strip()}"
                
        return entrada, corpo_autores, entrada_por_titulo

    def buscar_na_tabela_cutter(texto_para_busca, titulo_obra):
        if not texto_para_busca or not titulo_obra: return "X000x"
        url_csv = "https://raw.githubusercontent.com/Biblio-Khan/gerador-ficha-cat/refs/heads/main/cutter.csv"
        try:
            df = pd.read_csv(url_csv, sep=',', encoding='utf-8', quotechar='"')
        except Exception:
            return f"{texto_para_busca.strip().upper()[0]}200{titulo_obra.strip().lower()[0]}"
        
        df.columns = df.columns.str.strip().str.lower()
        col_nome = 'name' if 'name' in df.columns else df.columns[0]
        col_id = 'id' if 'id' in df.columns else df.columns[1]
        
        df['Name_Clean'] = df[col_nome].astype(str).str.strip().str.upper()
        sub_busca = texto_para_busca.strip().upper()
        
        match = df[df['Name_Clean'] <= sub_busca].sort_values(by='Name_Clean').tail(1)
        num = "200"
        if not match.empty:
            num = str(match[col_id].values[0]).strip().split('.')[0]
            
        titulo_limpo = titulo_obra.strip().upper()
        artigos = ["O ", "A ", "OS ", "AS ", "UM ", "UMA ", "UNS ", "UMAS "]
        for artigo in artigos:
            if titulo_limpo.startswith(artigo):
                titulo_limpo = titulo_limpo[len(artigo):].strip()
                break
        letra_titulo = titulo_limpo[0].lower() if titulo_limpo else "t"
        return f"{sub_busca[0]}{num}{letra_titulo}"

    def calcular_cutter(tipo_autor, autores_lista, entidade="", titulo="", tem_organizador=False, organizador_nome=""):
        if tipo_autor == "Entidade (Órgão/Instituição)" and entidade:
            texto_base = entidade
        elif tipo_autor == "Pessoa Física" and autores_lista and any(a.strip() for a in autores_lista):
            autor_principal = [a.strip() for a in autores_lista if a.strip()][0]
            partes = autor_principal.split()
            texto_base = partes[-1] if len(partes) > 1 else autor_principal
        elif tem_organizador or tipo_autor == "Organizador":
            partes_org = organizador_nome.strip().split()
            texto_base = partes_org[-1] if len(partes_org) > 1 else organizador_nome
        else:
            texto_base = "Autor"
        return buscar_na_tabela_cutter(texto_base, titulo)

    def gerar_marc21_completo(dados):
        marc_lines = [
            "000 00000nam a2200000 i 4500",
            f"100 1#$a{dados.get('entrada', '')}",
            f"245 10$a{dados.get('titulo', '')}"
        ]
    
        # Tag 260: Só adiciona se houver pelo menos a cidade ou a editora
        local = dados.get('local_editora', '')
        if local and local != " : ": # Verifica se não está vazio
            marc_lines.append(f"260 ##$a{local}")
    
        # Tag 300: Só adiciona se houver páginas OU dimensões
        paginas = dados.get('paginas', '')
        dimensoes = dados.get('dimensoes', '')
        if paginas or dimensoes:
            marc_lines.append(f"300 ##$a{paginas} p. ; {dimensoes} cm.")
    
        # Tag 502
        tipo = dados.get('tipo', '')
        if tipo and "Livro" not in tipo:
            inst = dados.get('instituicao', '')
            ano = dados.get('ano', '')
            marc_lines.append(f"502 ##$a{tipo} - {inst}, {ano}.")
    
        # Tags 650
        for assunto in dados.get('assuntos', []):
            if assunto:
                marc_lines.append(f"650 #4$a{assunto}")
    
        if dados.get('area'):
            marc_lines.append(f"650 #4$a{dados.get('area')}")

        return "\n".join(marc_lines)

   

    # =========================================================================
    # SISTEMA DE ABAS (CATALOGAÇÃO & CRÉDITOS LIMITADOS ATÉ 300)
    # =========================================================================
    tab_gerador, tab_financeiro, tab_produtividade = st.tabs([
    "Gerar Ficha", 
    "💳 Compra e Gestão de Créditos",
    "📊 Painel de Produtividade"
])

    with tab_gerador:
        if st.session_state.get("creditos_ativos", 0) <= 0:
            st.warning("🔒 O painel de salvamento está bloqueado. Adquira créditos ou aguarde a restauração para continuar.")

        st.title("Assistente de Catalogação — BiblioKhan")
        st.caption("Ficha Catalográfica, MARC21 e Referências ABNT.")
        
        st.markdown("---")
        container_lote = st.container()
        with container_lote:
            col_lote_1, col_lote_2, col_lote_3, col_lote_4 = st.columns([2, 1, 1, 1])
            qtd_fichas = len(st.session_state.lote_fichas)
            col_lote_1.subheader(f"📦 Lote: {qtd_fichas} Ficha(s)")
            
            if qtd_fichas > 0:
                # === COLUNA 2: DOCUMENTOS EM WORD ===
                # 1. Fichas Catalográficas (Word)
                arquivo_word_fichas = gerar_docx_lote([f["texto_ficha"] for f in st.session_state.lote_fichas])
                col_lote_2.download_button(
                    label="📄 Fichas (Word)",
                    data=arquivo_word_fichas,
                    file_name="lote_fichas.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # 2. Referências ABNT (Word) - NOVO BOTÃO
                arquivo_word_referencias = gerar_docx_referencias_lote(st.session_state.lote_fichas)
                col_lote_2.download_button(
                    label="Referências ABNT (Word)",
                    data=arquivo_word_referencias,
                    file_name="lote_referencias_abnt.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                # === COLUNA 3: REGISTROS MARC 21 ===
                conteudo_marc = "\n\n".join([gerar_marc21_completo(f["dados_marc"]) for f in st.session_state.lote_fichas])

                col_lote_3.download_button(
                    label="📥 MARC 21 (.mrc)",
                    data=conteudo_marc,
                    file_name="lote_juridico.mrc",
                    mime="text/plain"
                )

                col_lote_3.download_button(
                    label="MARC 21 (.txt)",
                    data=conteudo_marc,
                    file_name="lote_juridico.txt",
                    mime="text/plain"
                )
                
                # === COLUNA 4: AÇÕES ===
                if col_lote_4.button("🗑️ Limpar Lote"):
                    st.session_state.lote_fichas = []
                    st.rerun()
            else:
                col_lote_2.info("O lote está vazio.")
        
        st.markdown("---")
        fid = st.session_state.form_id
        col_esquerda, col_direita = st.columns(2)

        with col_esquerda:
            st.subheader("1. Metadados & Responsabilidade")
            classificacao = st.text_input("Número de Classificação (CDD ou CDU)", value="340.1", key=f"classificacao_{fid}")
            tipo_autor = st.radio("Tipo de Autoria Principal", ["Pessoa Física", "Entidade (Órgão/Instituição)"], horizontal=True, key=f"tipo_autor_{fid}")
            
            autores_lista = []
            entidade_nome = ""
            
            if tipo_autor == "Pessoa Física":
                qtd_autores_input = st.number_input("Quantidade de autores principais (0 se houver apenas Organizador)", min_value=0, max_value=10, value=1, key=f"qtd_autores_{fid}")
                for i in range(int(qtd_autores_input)):
                    autores_lista.append(st.text_input(f"Autor {i+1} (Nome Sobrenome)", key=f"autor_{fid}_{i}"))
            else:
                entidade_nome = st.text_input("Nome da Entidade (Ex: Brasil. Supremo Tribunal Federal)", key=f"entidade_nome_{fid}")
                
            titulo = st.text_input("Título Principal", key=f"titulo_{fid}")
            st.markdown("---")
            col_resp_1, col_resp_2 = st.columns(2)
            
            with col_resp_1:
                tem_organizador = st.checkbox("Possui Organizador/Coordenador?", key=f"tem_organizador_{fid}")
                organizador_nome = ""
                tipo_org, abreviatura_org = "", ""
                if tem_organizador:
                    papel = st.selectbox("Função:", ["Organizador", "Coordenador", "Compilador"], key=f"papel_{fid}")
                    organizador_nome = st.text_input("Nome do Responsável", key=f"organizador_nome_{fid}")
                    if papel == "Organizador": tipo_org, abreviatura_org = "organizado", "org."
                    elif papel == "Coordenador": tipo_org, abreviatura_org = "coordenado", "coord."
                    else: tipo_org, abreviatura_org = "compilado", "comp."
                    
            with col_resp_2:
                tem_tradutor = st.checkbox("A obra possui Tradutor?", key=f"tem_tradutor_{fid}")
                tradutor_nome = ""
                if tem_tradutor:
                    tradutor_nome = st.text_input("Nome do Tradutor (Nome Sobrenome)", key=f"trad_nome_{fid}")
            
            st.markdown("---")
            st.subheader("2. Publicação & Descrição Física")
            edicao = st.text_input("Edição e Volume (Ex: 2. ed., 3. ed. rev. e ampl.)", value="1. ed.", key=f"edicao_{fid}")
            editora = st.text_input("Editora", key=f"editora_{fid}")
            cidade = st.text_input("Cidade de Publicação", value="Brasília", key=f"cidade_{fid}")
            ano = st.text_input("Ano de Publicação", value="2026", key=f"ano_{fid}")
            paginas_input = st.text_input("Número de Páginas/Folhas", value="180", key=f"paginas_{fid}")
            dimensoes_input = st.text_input("Dimensões", value="30 cm", key=f"dimensoes_{fid}")
            
            tem_colecao = st.checkbox("Esta obra faz parte de uma Coleção / Série?", key=f"tem_colecao_{fid}")
            colecao_nome = ""
            if tem_colecao:
                colecao_nome = st.text_input("Nome da Coleção e Volume (Ex: Biblioteca jurídica, v. 12)", key=f"colecao_{fid}")

            st.markdown("---")
            st.subheader("3. Tipo de Documento / Trabalho Acadêmico")

            grau_academico = st.selectbox(
                "Tipo de Obra:", 
                ["Livro / Código / Obra Geral", "Tese (Doutorado)", "Dissertação (Mestrado)", "Monografia (Especialização)", "Monografia (Graduação)"],
                key=f"grau_academico_{fid}"
            )

            instituicao = ""
            area_concentracao = ""

            if grau_academico != "Livro / Código / Obra Geral":
                instituicao = st.text_input("Instituição / Universidade (Ex: Faculdade de Direito da USP):", key=f"instituicao_{fid}")
                area_concentracao = st.text_input("Área de Concentração / Curso (Ex: Direito Civil):", key=f"area_{fid}")
        
            isbn = st.text_input("ISBN (Ex: 978-65-0000-00-0)", key=f"isbn_{fid}")
            suporte = st.radio("Suporte da Obra", ["Impresso", "Digital"], horizontal=True, key=f"suporte_{fid}")
            url_acesso = st.text_input("URL de Acesso / DOI", key=f"url_acesso_{fid}") if suporte == "Digital" else ""

        with col_direita:
            st.subheader("3. Indexação por Assunto")
            st.markdown("##### 🏛️ Buscar no VCB do Senado Federal")
            termo_busca = st.text_input("Digite um termo para pesquisar:", key=f"termo_busca_{fid}")
            
            if termo_busca:
                resultados_vcb = buscar_vcb_senado(termo_busca)
                if resultados_vcb:
                    st.success(f"{len(resultados_vcb)} conceitos localizados no Senado!")
                    mapeamento_opcoes = {item["termo"]: item for item in resultados_vcb}
                    lista_opcoes = sorted(list(mapeamento_opcoes.keys()))
                    termo_selecionado = st.selectbox("Selecione o conceito oficial:", lista_opcoes, key=f"termo_sel_{fid}")
            
                    if st.button("➕ Vincular Assunto do Senado"):
                        if termo_selecionado not in st.session_state.assuntos_selecionados:
                            st.session_state.assuntos_selecionados.append(termo_selecionado)
                            st.rerun()
                else:
                    st.warning("Nenhum termo correspondente retornado pela API do Senado.")

            st.markdown("##### ✍️ Adicionar Assunto Manualmente")
            assunto_manual = st.text_input("Digite um assunto customizado:", key=f"assunto_manual_{fid}")
            if st.button("➕ Vincular Assunto Manual"):
                if assunto_manual.strip():
                    termo_limpo = assunto_manual.strip()
                    if termo_limpo not in st.session_state.assuntos_selecionados:
                        st.session_state.assuntos_selecionados.append(termo_limpo)
                        st.rerun()

            if st.session_state.assuntos_selecionados:
                st.write("**Assuntos Vinculados à Ficha:**")
        
                assunto_para_remover = None
                for idx, ass in enumerate(st.session_state.assuntos_selecionados):
                    col_assunto, col_excluir = st.columns([9, 1])
                    with col_assunto:
                        st.write(f"{idx+1}. {ass}")
                    with col_excluir:
                        if st.button("❌", key=f"remover_assunto_{idx}", help="Remover apenas este assunto"):
                            assunto_para_remover = ass
                    
                if assunto_para_remover:
                    st.session_state.assuntos_selecionados.remove(assunto_para_remover)
                    st.rerun()
                            
                # Mantemos o botão de limpar tudo, caso o usuário queira zerar a lista
                if st.button("🗑️ Limpar Todos os Assuntos"):
                    st.session_state.assuntos_selecionados = []
                    st.rerun()

            st.markdown("---")
            st.subheader("4. Fechamento e Visualização da Ficha")
                
            
            entrada_principal, responsabilidade, entrada_por_titulo = formatar_entrada_e_corpo(
                tipo_autor=tipo_autor, autores_lista=autores_lista, entidade=entidade_nome, titulo=titulo, 
                tem_organizador=tem_organizador, organizador_nome=organizador_nome, tipo_org=tipo_org, 
                tem_tradutor=tem_tradutor, tradutor_nome=tradutor_nome
            )
            
            cutter = calcular_cutter(tipo_autor, autores_lista, entidade=entidade_nome, titulo=titulo, tem_organizador=tem_organizador, organizador_nome=organizador_nome)
            dgm = " [recurso eletrônico]" if suporte == "Digital" else ""
            desc_fisica = f"1 recurso online ({paginas_input} p.) "if suporte == "Digital" else f"{paginas_input} p"
            if suporte != "Digital" and dimensoes_input.strip():
                desc_fisica = f"{desc_fisica} ; {dimensoes_input.strip()}"

            bloco_colecao = ""
            if tem_colecao and colecao_nome.strip():
                text_colecao = colecao_nome.strip()
                text_colecao = text_colecao[0].upper() + text_colecao[1:]
                bloco_colecao = f" ({text_colecao})"
                
            # === NOVO BLOCO: Gera a nota de trabalho acadêmico seguindo a ABNT ===
            nota_trabalho_str = ""
            if grau_academico != "Livro / Código / Obra Geral":
                inst_str = f" – {instituicao.strip()}" if instituicao.strip() else ""
                area_str = f" em {area_concentracao.strip()}" if area_concentracao.strip() else ""
                nota_trabalho_str = f"\n            {grau_academico}{area_str}{inst_str}, {ano.strip()}."

            nota_acesso = f"\n            Modo de acesso: {url_acesso}" if suporte == "Digital" and url_acesso else ""
            isbn_bloco = f"\n            ISBN {isbn}" if isbn.strip() else ""
            nota_traducao = f"\n            Traduzido de obra original." if tem_tradutor and tradutor_nome.strip() else ""
            ed_bloco = f"{edicao.strip()} – " if edicao.strip() else ""
            pub_bloco = f"{cidade.strip()} : {editora.strip()}, {ano.strip()}."
            
            string_assuntos = " ".join([f"{i+1}. {ass}" for i, ass in enumerate(st.session_state.assuntos_selecionados)])
            rastreabilidade = ""
            romanos = ["I", "II", "III", "IV", "V"]
            r_idx = 0
            
            if not entrada_por_titulo:
                rastreabilidade += f" {romanos[r_idx]}. Título."
                r_idx += 1
                
            if tem_organizador and organizador_nome.strip():
                partes_org = organizador_nome.strip().split()
                nome_invertido_org = f"{partes_org[-1].upper()}, {' '.join(partes_org[:-1])}" if len(partes_org) > 1 else organizador_nome.strip().upper()
                rastreabilidade += f" {romanos[r_idx]}. {nome_invertido_org}, {abreviatura_org}."
                r_idx += 1
                
            if tem_tradutor and tradutor_nome.strip():
                partes_trad = tradutor_nome.strip().split()
                nome_invertido_trad = f"{partes_trad[-1].upper()}, {' '.join(partes_trad[:-1])}" if len(partes_trad) > 1 else tradutor_nome.strip().upper()
                rastreabilidade += f" {romanos[r_idx]}. {nome_invertido_trad}, trad."
                r_idx += 1

            # === ADICIONADO {nota_trabalho_str} NAS DUAS STRINGS DA FICHA ===
            if entrada_por_titulo:
                txt_ficha = f"""{classificacao}
{cutter}   {titulo.strip()}{dgm} / {responsabilidade}. – {ed_bloco}{pub_bloco}
            {desc_fisica}.{bloco_colecao}{nota_trabalho_str}{nota_traducao}{nota_acesso}{isbn_bloco}
            
            {string_assuntos}{rastreabilidade}"""
            else:
                txt_ficha = f"""{classificacao}
{cutter}   {entrada_principal}
            {titulo.strip()}{dgm} / {responsabilidade}. – {ed_bloco}{pub_bloco}
            {desc_fisica}.{bloco_colecao}{nota_trabalho_str}{nota_traducao}{nota_acesso}{isbn_bloco}
            
            {string_assuntos}{rastreabilidade}"""

        
            st.text_area("Visualização Normativa (Fonte Monoespaçada)", value=txt_ficha, height=240)
    
 
            if st.button("💾 CONCLUIR FICHA E ENVIAR AO LOTE", disabled=st.session_state["creditos_ativos"] <= 0):
                valido = True
                if tipo_autor == "Pessoa Física" and not any(a.strip() for a in autores_lista) and not tem_organizador:
                    valido = False
                    st.error("⚠️ Informe ao menos um autor principal ou marque a opção de Organizador.")
                if not titulo.strip():
                    valido = False
                    st.error("⚠️ O Título Principal é obrigatório.")
    
                if valido:
                    with st.spinner("Gravando ficha e atualizando saldo na nuvem..."):
                        try:
                            uid = st.session_state["user_uid"]
                            email_usuario = st.session_state["usuario_atual"]
                            titulo_livro = titulo if titulo else "Não Informado"
                            lista_assuntos = st.session_state.get("assuntos_selecionados", [])

                            # 1. Desconta o crédito no Firestore
                            novo_saldo = descontar_credito_usuario(uid)
                
                            if novo_saldo is not None:
                                # 2. Salva o histórico na coleção 'historico' do Firestore
                                salvar_historico_firestore(
                                    uid_usuario=uid,
                                    email_usuario=email_usuario,
                                    titulo_livro=titulo_livro,
                                    assuntos_lista=lista_assuntos
                                )

                                # 3. Monta a ficha para o lote local
                                ficha_completa = {
                                    "texto_ficha": txt_ficha,
                                    "citacao_abnt": gerar_citacao_abnt_nbr6023(
                                        tipo_autor, autores_lista, entidade_nome, titulo,
                                        tem_organizador, organizador_nome, abreviatura_org,
                                        edicao, editora, cidade, ano, paginas_input,
                                        grau_academico, instituicao, area_concentracao, url_acesso

                                    ),
                                    "dados_marc": {
                                        "entrada": entrada_principal,
                                        "titulo": titulo,
                                        "local_editora": f"{cidade.strip()} : {editora.strip()}",
                                        "tipo": grau_academico,
                                        "instituicao": instituicao.strip(),
                                        "area": area_concentracao.strip(),
                                        "assuntos": lista_assuntos,
                                        "ano": ano.strip(),
                                        "paginas": paginas_input,
                                        "dimensoes": dimensoes_input
                                    }
                                }

                                st.session_state.lote_fichas.append(ficha_completa)
                                st.session_state["creditos_ativos"] = novo_saldo
                                st.session_state.form_id += 1
                                st.session_state.assuntos_selecionados = []
                                    
                                st.success(f"✅ Ficha guardada com sucesso! Créditos restantes: {novo_saldo}")
                                st.rerun()
                            else:
                                st.error("❌ Erro ao descontar crédito no banco de dados.")
                
                        except Exception as e:
                            st.error(f"❌ Erro ao processar requisição: {e}")  
                
# Abaixo, fora de qualquer bloco 'if' ou 'try', começa o tab_financeiro
    with tab_financeiro:
        st.header("💳 Gestão Financeira e Saldo")
        # ... resto do seu código da aba
        col_f1, col_f2 = st.columns(2)
    
    with col_f1:
        st.subheader("🔄 Sincronização")
        st.info(f"Seu sistema está vinculado ao e-mail: **{st.session_state['usuario_atual']}**")
        if st.button("Atualizar meu Saldo"):
            with st.spinner("Puxando dados atualizados do Sheets..."):
                atualizar_saldo_usuario(st.session_state["usuario_atual"])
                st.success("Saldo checado com sucesso!")
                st.rerun()

        with col_f2:
            st.subheader("🛒 Tabela de Preços")
            st.markdown("""
            * **20 Fichas** — R$ 55,00 
            * **30 Fichas** — R$ 80,00 
            * **100 Fichas** — R$ 240,00 
            * **300 Fichas** — R$ 660,00 
            * **600 Fichas** — R$ 1,200.00 
            """)
            st.info("🔑 **PIX:** `bibliokhancontato@gmail.com`")

        st.markdown("---")
        st.subheader("📩 Envio de Comprovante")
        
        with st.form("pix_form_original"):
            email_cliente = st.text_input("E-mail de Cadastro no Sistema", value=st.session_state["usuario_atual"], disabled=True)
           
            pacote_escolhido = st.selectbox(
                "Qual pacote de créditos você comprou?",
                options=[
                    "20 Fichas (R$ 55,00)",
                    "30 Fichas (R$ 80,00)",
                    "100 Fichas (R$ 240,00)",
                    "300 Fichas (R$ 660,00)",
                    "600 Fichas (R$ 1,200.00)"
                ]
            )
            
            comprovante = st.file_uploader("Anexe a imagem ou PDF do comprovante do PIX", type=["jpg", "png", "jpeg", "pdf"])
            
            if st.form_submit_button("Enviar para Restauração de Saldo"):
                if comprovante is not None:
                    with st.spinner("Enviando comprovante para o suporte... Por favor, aguarde."):
                        try:
                            tg_token = st.secrets["TELEGRAM_BOT_TOKEN"]
                            tg_chat = st.secrets["TELEGRAM_CHAT_ID"]
    
                            fuso_brasilia = timezone(timedelta(hours=-3))
                            data_hora_br = datetime.now(fuso_brasilia).strftime('%d/%m/%Y %H:%M:%S')
                            
                            texto_notificacao = (
                                f"🔥 *NOVO COMPROVANTE RECEBIDO!*\n\n"
                                f"📧 *E-mail do Cliente:* {st.session_state['usuario_atual']}\n"
                                f"💰 *Pacote Escolhido:* {pacote_escolhido}\n"
                                f"📅 *Data/Hora:* {data_hora_br}"
                            )
                            
                            url_api_telegram = f"https://api.telegram.org/bot{tg_token}/sendPhoto"
                            ficheiro_envio = {"photo": (comprovante.name, comprovante.getvalue(), comprovante.type)}
                            dados_requisicao = {"chat_id": tg_chat, "caption": texto_notificacao, "parse_mode": "Markdown"}
                            
                            resposta_tg = requests.post(url_api_telegram, data=dados_requisicao, files=ficheiro_envio, timeout=15)
                            
                            if resposta_tg.status_code == 200:
                                st.success("✅ Comprovante enviado com sucesso!")
                                st.info("⏳ O seu saldo será atualizado assim que a validação for concluída.")
                            else:
                                st.error(f"Erro na API de comunicação (Código {resposta_tg.status_code}).")
                        except Exception as e:
                            st.error(f"Erro ao disparar arquivo de envio: {e}")
                else:
                    st.error("❌ Por favor, informe o seu nome completo e anexe o arquivo do comprovante.")

# ---------------------------------------------------------------------
# NOVA ABA: PAINEL DE PRODUTIVIDADE JURÍDICA (COM TRAVA DE LOGIN)
# ---------------------------------------------------------------------
# Só executa este bloco se o usuário já tiver passado pela tela de login
if st.session_state.get("usuario_atual"):

    # Verifica dinamicamente se a aba foi criada no topo do arquivo
    if 'tab_produtividade' not in locals() and 'tab_produtividade' not in globals():
        st.markdown("---")
        tab_produtividade = st.container()

    with tab_produtividade:
        st.title("Painel de Produtividade")
        st.subheader(f"Análise de Obras Processadas por {st.session_state.get('usuario_atual', 'Usuário')}")

        with st.spinner("Carregando dados de produtividade..."):
            # Usando o user_uid para buscar corretamente no Firestore
            uid_atual = st.session_state.get("user_uid", "")
            dados = api_obter_produtividade_juridica(uid_atual)

        import pandas as pd
        
        # Verifica se o objeto 'dados' é um DataFrame válido e não está vazio
        if isinstance(dados, pd.DataFrame) and not dados.empty:
            st.write(f"Total de registros encontrados: {len(dados)}")
            
            # 1. Converte os dados recebidos para um DataFrame do Pandas
            df = pd.DataFrame(dados)

            # 2. Coleta todos os assuntos, quebra pelas vírgulas e limpa os espaços
            todos_assuntos = []
            for linha_assunto in df['assunto']:
                if linha_assunto: 
                    if str(linha_assunto) != "Não informado":
                        partes = [a.strip().title() for a in str(linha_assunto).split(",") if a.strip()]
                        todos_assuntos.extend(partes)

            # 3. Conta a frequência de cada assunto individual
            if todos_assuntos:
                df_contagem = pd.DataFrame(todos_assuntos, columns=["Área/Assunto"]).value_counts().reset_index(name="Quantidade")
            else:
                df_contagem = pd.DataFrame()

            # 4. Mostra os cartões de resumo (Métricas)
            col_card1, col_card2 = st.columns(2)
            with col_card1:
                st.metric("Total de Processos/Livros", len(df))
            with col_card2:
                st.metric("Total de Assuntos Mapeados", len(df_contagem))

            st.markdown("---")
            
            # 5. Renderiza o Gráfico de Barras se houver assuntos mapeados
            if not df_contagem.empty:
                st.write("### Temas mais Demandados nas suas Fichas")
                st.bar_chart(
                    data=df_contagem,
                    x="Área/Assunto",
                    y="Quantidade",
                    color="#0077B6", 
                    use_container_width=True
                )
                st.markdown("---")

            # 6. Histórico de Obras Processadas e Opção de Download
            st.write("### Histórico de Fichas Emitidas")
            
            df_exibicao = df.copy()
            
            df_exibicao = df_exibicao.rename(columns={
                "data": "Data/Hora",
                "titulo": "Título da Obra",
                "assunto": "Assuntos Indexados"
            })
            
            if "Data/Hora" in df_exibicao.columns:
                try:
                    df_exibicao["Data/Hora"] = pd.to_datetime(df_exibicao["Data/Hora"]).dt.strftime('%d/%m/%Y %H:%M')
                except:
                    pass 

            colunas_relatorio = ["Data/Hora", "Título da Obra", "Assuntos Indexados"]
            df_final = df_exibicao[colunas_relatorio]

            # === BOTÃO DE DOWNLOAD ===
            csv_dados = df_final.to_csv(index=False, sep=";").encode('utf-8-sig')
            
            st.download_button(
                label="📥 Baixar Relatório em CSV (Excel)",
                data=csv_dados,
                file_name=f"produtividade_juridica_{st.session_state['usuario_atual'].split('@')[0]}.csv",
                mime="text/csv",
                use_container_width=True
            )
            
            st.write("") 
            
            st.dataframe(
                df_final, 
                use_container_width=True,
                hide_index=True
            )
        else:
            st.info("Você ainda não possui registros de fichas geradas.")
